import os
import fitz  # PyMuPDF
from docx import Document
from llama_index.core import Document as LlamaDocument
from llama_index.core.node_parser import SentenceSplitter

UPLOAD_DIR = "uploads"

def extract_text_from_pdf(filepath: str) -> str:
    text = ""
    doc = fitz.open(filepath)
    for page in doc:
        text += page.get_text()
    return text

def extract_text_from_docx(filepath: str) -> str:
    doc = Document(filepath)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def extract_text_from_txt(filepath: str) -> str:
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def extract_text(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    elif ext == ".txt":
        return extract_text_from_txt(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def chunk_text(text: str, filename: str) -> list:
    doc = LlamaDocument(
        text=text,
        metadata={"source": filename}
    )
    splitter = SentenceSplitter(
        chunk_size=512,
        chunk_overlap=64
    )
    nodes = splitter.get_nodes_from_documents([doc])
    print(f"[Ingestor] Created {len(nodes)} chunks from {filename}")
    return nodes

def ingest_file(filename: str) -> list:
    filepath = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    
    print(f"[Ingestor] Extracting text from {filename}...")
    text = extract_text(filepath)
    
    if not text.strip():
        raise ValueError(f"No text could be extracted from {filename}")
    
    print(f"[Ingestor] Extracted {len(text)} characters")
    nodes = chunk_text(text, filename)
    return nodes